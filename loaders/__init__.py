"""Loader — 文档加载模块

职责：读文件 -> 提取纯文本
隔离边界：输入文件路径 -> 输出 Document 列表
"""
import logging
from abc import ABC, abstractmethod
from pathlib import Path
from typing import List, Optional

from models import Document

logger = logging.getLogger(__name__)


class Loader(ABC):
    """文档加载抽象基类"""

    @abstractmethod
    def load(self, path: str) -> List[Document]:
        ...

    @staticmethod
    def detect(path: str) -> "Loader":
        ext = Path(path).suffix.lower()
        if ext == ".epub":
            return EpubLoader()
        elif ext == ".pdf":
            return PdfLoader()
        elif ext == ".docx":
            return DocxLoader()
        raise ValueError(f"Unsupported file format: {ext}")


class EpubLoader(Loader):
    """epub 加载，逐章保留章节结构"""

    def load(self, path: str) -> List[Document]:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup

        book = epub.read_epub(path)
        docs: List[Document] = []

        for item in book.get_items():
            if item.get_type() != ebooklib.ITEM_DOCUMENT:
                continue

            name = item.get_name()
            content = item.get_content()

            # 检测编码（借鉴 Open WebUI CJK 检测）
            html_text = self._decode(content)
            if not html_text.strip():
                continue

            # 解析 HTML 提取文本 + 保留标题结构
            soup = BeautifulSoup(html_text, "html.parser")

            # 提取标题（优先 <title> 或第一个 <h1>）
            title = ""
            title_tag = soup.find("title")
            if title_tag and title_tag.get_text(strip=True):
                title = title_tag.get_text(strip=True)
            if not title:
                h1 = soup.find("h1")
                if h1 and h1.get_text(strip=True):
                    title = h1.get_text(strip=True)

            # 提取正文：h1~h6 转为 md 标题，去掉标签
            lines = []
            for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p",
                                       "pre", "blockquote", "li"]):
                text = tag.get_text(strip=True)
                if not text:
                    continue
                level = int(tag.name[1]) if tag.name.startswith("h") else 0
                if level > 0:
                    lines.append(f"{'#' * level} {text}")
                elif tag.name == "li":
                    lines.append(f"- {text}")
                else:
                    lines.append(text)

            text_content = "\n\n".join(lines)
            if not text_content.strip():
                continue

            docs.append(Document(
                text=text_content,
                metadata={
                    "source": Path(path).name,
                    "file_path": path,
                    "chapter": name,
                    "title": title,
                },
            ))

        return docs

    @staticmethod
    def _decode(content: bytes) -> str:
        """尝试多种解码，支持 GB18030 等 CJK 编码"""
        for enc in ("utf-8", "gb18030", "gbk", "gb2312", "big5", "shift-jis"):
            try:
                return content.decode(enc)
            except (UnicodeDecodeError, LookupError):
                continue
        # fallback: 忽略错误
        return content.decode("utf-8", errors="ignore")


class PdfLoader(Loader):
    """PDF 加载（PyMuPDF）"""

    def load(self, path: str) -> List[Document]:
        import fitz  # PyMuPDF

        doc = fitz.open(path)
        docs: List[Document] = []
        source = Path(path).name

        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text().strip()
            if not text:
                continue
            docs.append(Document(
                text=text,
                metadata={
                    "source": source,
                    "file_path": path,
                    "page": page_num + 1,
                },
            ))

        doc.close()
        return docs


class DocxLoader(Loader):
    """Word 文档加载"""

    def load(self, path: str) -> List[Document]:
        from docx import Document as DocxDoc

        doc = DocxDoc(path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            return []
        return [Document(
            text="\n\n".join(paragraphs),
            metadata={
                "source": Path(path).name,
                "file_path": path,
            },
        )]
